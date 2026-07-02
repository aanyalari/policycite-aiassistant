"""Build the three-page Cotiviti assessment report."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("PolicyCite_RAG_Report_Aanya_Lari.docx")
BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(20, 32, 51)
MUTED = RGBColor(91, 102, 116)
LIGHT = "E8EEF5"


def set_font(run, size=11, bold=False, italic=False, color=DARK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def add_body(doc, text, *, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        body = p.add_run(text[len(bold_lead):])
        set_font(body)
    else:
        set_font(p.add_run(text))
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    set_font(p.add_run(text), size=10.5)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(
        header.add_run("POLICYCITE-RAG  |  COTIVITI INTERNSHIP ASSESSMENT"),
        size=8.5,
        bold=True,
        color=MUTED,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Aanya Lari  |  "), size=9, color=MUTED)
    add_field(footer, "PAGE")


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(2)
    set_font(title.add_run("PolicyCite-RAG"), size=24, bold=True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    set_font(
        subtitle.add_run("Verifiable Generative AI for Healthcare Payment Policy"),
        size=14,
        bold=True,
        color=DARK,
    )
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    set_font(
        meta.add_run("Topic 3: Content Management in Health Care  |  July 2026"),
        size=9.5,
        color=MUTED,
    )

    add_heading(doc, "Concept and industry trend")
    add_body(
        doc,
        "Healthcare payment operations depend on policies that define claim formats, filing limits, prior-authorization obligations, and payer-provider workflows. Generative AI can make this content easier to search and summarize, but a fluent answer with a document list is not necessarily verifiable. A page may be relevant to the question while failing to support a particular date, exclusion, or requirement. In payment-policy work, that gap can turn a convenient assistant into an audit and governance risk.",
    )
    add_body(
        doc,
        "Retrieval-augmented generation (RAG) reduces unsupported generation by supplying external documents to a language model. The emerging research trend is finer-grained attribution: each generated statement is connected to evidence that a reviewer can inspect. MedCite demonstrates this direction for biomedical question answering through RAG, statement-level re-retrieval, citation attribution, and separate citation-recall and citation-precision measures (Wang et al., 2025). PolicyCite-RAG adapts those ideas to public Centers for Medicare & Medicaid Services (CMS) payment-policy documents.",
    )

    add_heading(doc, "Proof of concept")
    add_body(
        doc,
        "The prototype preserves an inherited FastAPI and Streamlit RAG baseline using FAISS, BM25, and Reciprocal Rank Fusion. The new citation-assurance path records the exact passages supplied to answer generation, extracts factual statements, re-retrieves evidence for each statement, and uses a separately configured local model to issue a strict binary SUPPORTED or NOT_SUPPORTED judgment. Supported statements receive a validated document, display page, and exact excerpt; unsupported statements trigger a review recommendation. Generation context and post-generation evidence remain separate so later evidence is not misrepresented as original grounding.",
    )
    add_body(
        doc,
        "The assessment corpus contains three public CMS documents covering Medicare timely filing, CMS-1500/837P billing, and prior authorization. A frozen ten-question evaluation includes direct, multi-statement, and unsupported-premise cases. The design intentionally favors a small reproducible vertical slice over agents, model fine-tuning, or production workflow infrastructure.",
    )

    add_heading(doc, "Measured result")
    add_body(
        doc,
        "In the provenance-fixed run, both conditions evaluated the exact same generated answers. PolicyCite reduced attached citations from 43 to 18 and improved independently reviewed citation precision from 32/43 (74.4%) to 18/18 (100.0%). Coverage decreased from 19/21 (90.5%) to 16/21 (76.2%) because the stricter pipeline declined to attach evidence to five statements. Citation F1 increased from 81.7% to 86.5%, complete-answer retention remained 7/10, and median overhead was 3.69 seconds. This is a precision-coverage tradeoff, not a universal performance claim. Project-author label signoff remains required before external submission. Later answer-cleanup guardrails are unit-tested but must be rebenchmarked before these values are presented as final current-system results.",
    )

    doc.add_page_break()
    add_heading(doc, "Opportunities for Cotiviti")
    add_body(
        doc,
        "Policy citation assurance could serve as a reusable governance layer for content-heavy treatment, payment, and operations workflows. The immediate value is not autonomous policy interpretation; it is reducing reviewer effort while preserving an inspectable trail from generated language to authoritative text.",
    )
    add_bullet(doc, "Policy research assistant: answer narrow billing or payment questions while exposing the exact supporting passage and page.")
    add_bullet(doc, "Content-quality control: evaluate summaries, policy-change explanations, or coding guidance at statement level before publication.")
    add_bullet(doc, "Governance service: export statement, evidence, model verdict, reviewer decision, corpus version, and timestamp as an audit record.")
    add_bullet(doc, "Evaluation harness: compare retrievers, models, and prompts using frozen questions and human-reviewed evidence rather than model self-grading.")

    add_heading(doc, "Threats and limitations")
    add_body(
        doc,
        "A citation does not prove that a document is current, controlling, or correctly interpreted. Retrieval may miss the best passage; an attribution model may accept related but incomplete evidence; PDF extraction can damage tables or footnotes; and a small local generator can emit awkward or irrelevant text. The present binary verdict also hides partial support, while the three-document corpus and ten-question set cannot establish production performance. Human review, effective-date metadata, access controls, monitoring, and domain-expert validation remain necessary. CMS itself distinguishes precise operational requirements, such as filing windows and prior-authorization timeframes, whose qualifiers must be preserved (CMS, 2011, 2024).",
    )

    add_heading(doc, "Strategic recommendation")
    add_body(
        doc,
        "Cotiviti should explore a bounded policy-assurance pilot rather than a general-purpose healthcare chatbot. Start with one versioned policy domain, require evidence for every material statement, and route unsupported output to a qualified reviewer. Measure statement coverage, citation precision, answer completeness, reviewer agreement, latency, and time saved. A second experiment should compare the current RRF citation retrieval with the lexical-first, semantic-second approach highlighted by MedCite, because exact policy terminology may reward lexical precision.",
    )
    add_body(
        doc,
        "The recommended investment is therefore a reusable evidence and evaluation layer—not a claim that generated text can replace policy experts. If the pilot demonstrates high precision, acceptable reviewer agreement, and measurable time savings, Cotiviti could extend the pattern to policy-change comparison, structured rule extraction, and controlled internal knowledge workflows. This path aligns technical experimentation with AI governance and converts a research insight into a practical, auditable capability.",
    )

    add_heading(doc, "Decision")
    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(0)
    p_pr = callout._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    p_pr.append(shd)
    set_font(
        callout.add_run(
            "Proceed with a narrow, reviewer-in-the-loop policy citation pilot and treat precision, provenance, and evaluation discipline as the product—not decorative citations."
        ),
        size=11,
        bold=True,
        color=BLUE,
    )

    doc.add_page_break()
    add_heading(doc, "References")
    references = [
        "Centers for Medicare & Medicaid Services. (2011, January 21). Changes to the time limits for filing Medicare fee-for-service claims (Transmittal 2140, Change Request 7270). U.S. Department of Health and Human Services.",
        "Centers for Medicare & Medicaid Services. (2024, January 17). CMS interoperability and prior authorization final rule CMS-0057-F: Fact sheet. U.S. Department of Health and Human Services.",
        "Centers for Medicare & Medicaid Services. (2025, December). Medicare billing: CMS-1500 & 837P (MLN006976). U.S. Department of Health and Human Services.",
        "Cotiviti. (2026). Intern - Generative AI Research Engineer: Position description (ID 2026-19341).",
        "Cotiviti. (n.d.). Intern assessment directions.",
        "Wang, X., Tan, M., Jin, Q., Xiong, G., Hu, Y., Zhang, A., Lu, Z., & Zhang, M. (2025). MedCite: Can language models generate verifiable text for medicine? arXiv. https://arxiv.org/abs/2506.06605",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(9)
        p.paragraph_format.line_spacing = 1.1
        set_font(p.add_run(reference), size=11)

    add_heading(doc, "Prototype evidence")
    add_body(
        doc,
        "PolicyCite-RAG source code, frozen questions, per-question traces, and timestamped results are included in the accompanying GitHub repository. The original failed run and invalidated label run are retained as audit history.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
